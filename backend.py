from typing import Optional, List
import io
from docx import Document
import os
from tkinter import filedialog
import json
import difflib
import tempfile
from docx.enum.section import WD_SECTION_START
from docx.shared import Inches

from nda_reviewer.llms.base import LLMProtocol
from nda_reviewer.llms.openai import OpenAIHandler
from nda_reviewer.utils.redlines import apply_redline


class Backend:
    def __init__(self) -> None:
        self.temperature = 0.1
        self._llm = OpenAIHandler(self.temperature)
        self.nda_content = None
        self.guidelines = None
        self.revised_nda = None
        self.suggested_changes = None
        self.original_docx_path = None
        self.revised_docx_path = None
        self.recipient = None

    def get_stream_response(self, user_input: str):
        if self._llm is None:
            raise ValueError("No Language Model Has Been Selected.")
        return self._llm.stream_response(user_input)

    def export_conversation(self):
        if self._llm is None:
            return
        new_file = filedialog.asksaveasfilename(
            initialfile="Untitled.txt",
            defaultextension=".txt",
            filetypes=[("Text File", "*.txt")],
        )
        if not new_file:
            return
        with open(new_file, "w") as f:
            f.write(self._llm.export_conversation())

    def upload_nda(self):
        file_path = filedialog.askopenfilename(
            filetypes=[("Word Document", "*.docx"), ("Text File", "*.txt")]
        )
        if not file_path:
            return "No file selected"

        self.original_docx_path = file_path
        with open(file_path, "rb") as file:
            if file_path.endswith(".docx"):
                doc = Document(file)
                content = "\n".join([para.text for para in doc.paragraphs])
            else:
                content = file.read().decode("utf-8")

        self.nda_content = content
        return f"NDA uploaded successfully: {os.path.basename(file_path)}"

    def upload_guidelines(self):
        file_path = filedialog.askopenfilename(
            filetypes=[("Text File", "*.txt"), ("Word Document", "*.docx")]
        )
        if not file_path:
            return "No file selected"

        with open(file_path, "rb") as file:
            if file_path.endswith(".docx"):
                doc = Document(file)
                content = "\n".join([para.text for para in doc.paragraphs])
            else:
                content = file.read().decode("utf-8")

        self.guidelines = content
        return f"Guidelines uploaded successfully: {os.path.basename(file_path)}"

    def download_revised_nda(self):
        if not self.revised_docx_path:
            return "No revised NDA available. Please analyze and revise the NDA first."

        original_filename = os.path.basename(self.original_docx_path)
        suggested_filename = f"{os.path.splitext(original_filename)[0]} - LPC Redline.docx"

        file_path = filedialog.asksaveasfilename(
            initialfile=suggested_filename,
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")],
        )
        if not file_path:
            return "Download cancelled"

        try:
            result = self.generate_redlined_docx(file_path)
            return result
        except Exception as e:
            error_message = f"An error occurred while generating the redlined document: {str(e)}"
            print(error_message)  # For debugging
            return error_message

    def analyze_and_revise_nda(self):
        if not hasattr(self, "nda_content") or not hasattr(self, "guidelines"):
            return "Please upload both NDA and guidelines before analyzing."

        if self._llm is None:
            raise ValueError("No Language Model Has Been Selected.")

        from datetime import datetime

        current_date = datetime.now().strftime("%B %d, %Y")
        recipient = self.recipient if self.recipient else "Long Point Capital, LLC"

        system_prompt = f"""You are an AI assistant specialized in analyzing and revising Non-Disclosure Agreements (NDAs). 
        Your task is to review the provided NDA and suggest revisions based on the given guidelines. 
        Additionally, identify any placeholders or sections that need to be filled with specific information.
        Use the following information to fill in appropriate placeholders:
        - Current Date: {current_date}
        - Recipient/Potential Buyer: {recipient}
        IMPORTANT: Always use "{recipient}" as the Recipient/Potential Buyer when filling in placeholders or suggesting changes.
        Provide your analysis and revisions in a clear, structured format."""

        examples = [
            {
                "original": "The Receiving Party agrees that it will only share Confidential Information with those of its directors, officers, employees, debt financing sources, advisors, legal counselors, consultants, agents and other representatives who need to know such information for the sole purpose of assisting the Receiving Party in evaluating a Transaction (collectively, its 'Representatives').",
                "revised": "The Receiving Party agrees that it will only share Confidential Information with those of its directors, officers, employees, financing sources, advisors, legal counselors, consultants, agents and other representatives who need to know such information for the sole purpose of assisting the Receiving Party in evaluating a Transaction (collectively, its 'Representatives').",
                "justification": "Removed 'debt' before 'financing sources' to allow for broader sharing with various types of financing sources, not just debt financing.",
            },
            {
                "original": "At the request of the Disclosing Party, the Receiving Party will return or destroy all Confidential Information and shall cause its Representatives to do the same. The Receiving Party will certify such destruction in writing signed by an officer of the Receiving Party.",
                "revised": "At the request of the Disclosing Party, the Receiving Party will return or destroy all Confidential Information and shall cause its Representatives to do the same. The Receiving Party will certify such destruction in writing signed by an officer of the Receiving Party only if so requested in writing.",
                "justification": "Added condition for written certification of destruction only when specifically requested, reducing administrative burden on the Receiving Party.",
            },
            {
                "original": "[No clause addressing non-compete]",
                "revised": "Non-Compete. The Company recognizes that the Receiving Party or its Representatives may be or become engaged in the research, development, production, marketing and sale of services or products similar to the Company's and that these services or products may be competitive with those of the Company and may display the same or similar functionality. Therefore, nothing in this Agreement shall be construed to prevent the Receiving Party, its officers, its employees or its Representatives from engaging independently in such activities now or in the future.",
                "justification": "Added a non-compete clause that acknowledges the Receiving Party's potential involvement in similar businesses, ensuring the NDA doesn't restrict their normal business operations.",
            },
            {
                "original": "The Receiving Party covenants and agrees that for a period of two (2) years from the date of this Agreement, the Receiving Party and its affiliates will not, directly or indirectly, solicit, induce, encourage or attempt to solicit, induce or encourage any employee of the Disclosing Party or any of its affiliates or subsidiaries to terminate such person's employment or relationship with the Disclosing Party or such affiliate or subsidiary in order to become an employee, consultant or contractor, to or for the Receiving Party or its affiliates.",
                "revised": "The Receiving Party covenants and agrees that for a period of two (2) years from the date of this Agreement, the Receiving Party and its affiliates will not, directly or indirectly, solicit, induce, encourage or attempt to solicit, induce or encourage any employee of the Disclosing Party or any of its affiliates or subsidiaries to terminate such person's employment or relationship with the Disclosing Party or such affiliate or subsidiary in order to become an employee, consultant or contractor, to or for the Receiving Party or its affiliates; provided, however, that the foregoing provisions shall not preclude the Receiving Party and its affiliates from: (i) the use of public general advertisements or search firms (in each case, not directed at, or targeted to the Disclosing Party, its affiliates, its subsidiaries or their employees) or the hiring of any person who responds thereto, or (ii) the soliciting or hiring any such person who is not employed by the Company for at least six (6) months prior to the commencement of any such solicitation or employment discussions, or (iii) contact the Receiving Party or its Representatives on his or her own volition.",
                "justification": "Added exceptions to the non-solicitation clause to allow for general recruitment activities and hiring of individuals who independently reach out, providing more flexibility while still protecting the Disclosing Party's interests.",
            },
            {
                "original": "The Receiving Party acknowledges that none of the Disclosing Party, Guggenheim Securities or any of the Disclosing Party's other representatives (i) makes any express or implied representation or warranty as to the completeness or accuracy of the Confidential Information, or (ii) will have any liability whatsoever to the Receiving Party or any of its Representatives relating to or arising from their review or use of the Confidential Information.",
                "revised": "The Receiving Party acknowledges that none of the Disclosing Party, Guggenheim Securities or any of the Disclosing Party's other representatives (i) makes any express or implied representation or warranty as to the completeness or accuracy of the Confidential Information, or (ii) will have any liability whatsoever to the Receiving Party or any of its Representatives relating to or arising from their review or use of the Confidential Information, except as agreed to in definitive documentation.",
                "justification": "Added an exception for liability as agreed in definitive documentation, allowing for potential liability to be established in future agreements while maintaining general protection for the Disclosing Party.",
            },
        ]

        example_prompt = "Here are some examples of NDA revisions:\n\n"
        for example in examples:
            example_prompt += f"Original: {example['original']}\n"
            example_prompt += f"Revised: {example['revised']}\n"
            example_prompt += f"Justification: {example['justification']}\n\n"

        user_input = f"""Please analyze and revise the following NDA according to these guidelines:

        Guidelines:
        {self.guidelines}

        {example_prompt}

        NDA:
        {self.nda_content}

        Provide your analysis and revised NDA as a list of JSON objects, where each object represents a suggested change:
        [
            {{
                "original_text": "Full original paragraph or placeholder from the NDA",
                "suggested_change": "Full revised paragraph with suggested changes or filled-in information",
                "justification": "Explanation for why this change was suggested or what information was filled in"
            }},
            // ... more objects for other paragraphs or placeholders
        ]
        Include paragraphs that need changes based on the guidelines AND any placeholders or sections that need to be filled with specific information.
        """

        try:
            response = self._llm.analyze_documents(system_prompt, user_input)

            if isinstance(response, list) and len(response) > 0:
                if "raw_content" in response[0]:
                    return f"Analysis complete, but the response was not in the expected format. Here's the raw response:\n\n{response[0]['raw_content']}"
                else:
                    self.suggested_changes = response
                    return "Analysis complete. Ready to review changes."
            else:
                raise ValueError("Unexpected response format from the language model.")
        except Exception as e:
            return f"An error occurred during analysis: {str(e)}"

    def review_changes(self):
        if not hasattr(self, "suggested_changes"):
            return "No changes to review. Please analyze the NDA first."

        for change in self.suggested_changes:
            yield change

    def apply_approved_changes(self, approved_changes):
        if not self.original_docx_path:
            return "No original NDA found. Please upload an NDA first."

        # Load the original document
        doc = Document(self.original_docx_path)

        # Apply changes
        for change in approved_changes:
            for paragraph in doc.paragraphs:
                if change["original_text"] in paragraph.text:
                    paragraph.text = paragraph.text.replace(
                        change["original_text"], change["suggested_change"]
                    )

        # Save the revised document
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
            doc.save(tmp_file.name)
            self.revised_docx_path = tmp_file.name

        # Set the revised_nda attribute
        self.revised_nda = "\n".join([para.text for para in doc.paragraphs])

        return "Changes applied successfully. You can now download the revised and redlined NDA."

    def generate_redlined_docx(self, output_path):
        if not self.original_docx_path or not self.revised_docx_path:
            return "Original or revised NDA not found. Please upload and revise the NDA first."

        return apply_redline(self.original_docx_path, self.revised_docx_path, output_path)

    def set_system_prompt(self, prompt: str):
        if isinstance(self._llm, OpenAIHandler):
            self._llm.set_system_prompt(prompt)

    def send_message(self, user_input: str) -> str:
        if self._llm is None:
            raise ValueError("No Language Model Has Been Selected.")
        return self._llm.get_response(user_input)

    def clear_conversation(self):
        if isinstance(self._llm, OpenAIHandler):
            self._llm.clear_messages()

    def set_recipient(self, recipient: str):
        self.recipient = recipient
