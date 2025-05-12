from flask import Flask, request, jsonify, send_file, render_template
from docx import Document
from io import BytesIO
import logging
import os
import json
import base64
from dotenv import load_dotenv
from datetime import datetime
import hmac
import hashlib
import uuid
import openai
from packaging import version
from openai import OpenAI
from advanced_redliner import AdvancedDocumentProcessor  # Import the advanced redliner

# Load environment variables
load_dotenv()

# Set up logging
logging.basicConfig(
    level=logging.DEBUG, format="%(asctime)s - %(name)s - %(levelname)s - %(message)s"
)
logger = logging.getLogger(__name__)

# Initialize Flask app
application = Flask(__name__)
app = application  # Add this line to make app available for gunicorn

# Get OpenAI configuration
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
OPENAI_MODEL = "gpt-4"  # Using constant model name

if not OPENAI_API_KEY:
    raise ValueError("OPENAI_API_KEY environment variable is not set")

# Detect OpenAI version
openai_version = getattr(openai, "__version__", "0.0.0")

if version.parse(openai_version) >= version.parse("1.0.0"):
    # v1.x style
    client = OpenAI(api_key=OPENAI_API_KEY)
    openai_api_mode = "v1"
    logger.info(
        f"Using OpenAI client v1.x (version {getattr(client, '__version__', 'unknown')})"
    )
else:
    # v0.x style
    openai.api_key = OPENAI_API_KEY
    client = openai
    openai_api_mode = "v0"
    logger.info(f"Using OpenAI client v0.x (version {openai_version})")


class NDAProcessingError(Exception):
    """Custom exception for NDA processing errors"""

    pass


def apply_changes_to_document(doc: Document, changes: list) -> Document:
    """Apply changes directly to document paragraphs while preserving formatting"""
    changes_made = False

    for paragraph in doc.paragraphs:
        original_text = paragraph.text
        modified_text = original_text

        for change in changes:
            if change["original_text"] in modified_text:
                modified_text = modified_text.replace(
                    change["original_text"], change["new_text"]
                )
                changes_made = True
                logger.info(
                    f"Applied change: {change['original_text']} → {change['new_text']}"
                )

        if modified_text != original_text:
            # Clear the paragraph and add new text with formatting
            paragraph.clear()
            paragraph.add_run(modified_text)

    if not changes_made:
        logger.warning("No changes were applied to the document")

    return doc


def load_nda_checklist() -> str:
    """Load NDA Checklist text directly from file"""
    try:
        with open("documents/NDA Checklist.txt", "r") as f:
            checklist_content = f.read()
            logger.info("Loaded NDA Checklist content")
            return checklist_content
    except FileNotFoundError:
        logger.error("NDA Checklist file not found")
        return "NDA Checklist not found"
    except Exception as e:
        logger.error(f"Error loading NDA Checklist: {str(e)}")
        return "Error loading NDA Checklist"


def analyze_nda_with_openai(nda_text: str) -> dict:
    """Analyze NDA text using OpenAI"""
    try:
        # Load checklist and prepare prompt
        checklist_content = load_nda_checklist()

        # Get current date in standard legal format
        current_date = datetime.now().strftime("%B %d, %Y")

        prompt = f"""
        EXACT TEXT MATCHING NDA REVIEW

        This is a specialized NDA review task where EXACT text matching is CRITICAL. The document processing system can ONLY find and replace text that matches EXACTLY character-for-character.

        NDA CHECKLIST:
        {checklist_content}

        NDA TEXT:
        {nda_text}

        YOUR TASK:
        1. BE COMPREHENSIVE - Find ALL sections in the NDA that could benefit from changes
        2. Identify EVERY chance to improve the document based on our checklist
        3. COPY-PASTE the EXACT text as it appears in the document
        4. Create targeted replacement text that incorporates our preferred language
        5. Fill in any blank dates with "{current_date}" and any blank buyer/party names with "Long Point Capital"

        ADDITIONAL INSTRUCTIONS:
        - Look for blank fields in the document (often marked with ___________ or left empty)
        - Search for patterns like "dated as of _______" or "entered into as of _______" and replace with the current date
        - Search for patterns like "and ____________ (the "Buyer")" or similar fields for the second party
        - Make sure to copy the EXACT text with all surrounding context - do not guess or approximate

        CRITICAL REQUIREMENTS:
        - The original_text MUST be an EXACT COPY from the document with no changes
        - PRESERVE ALL whitespace, line breaks, and punctuation exactly as they appear
        - Even a single character difference will cause the replacement to fail
        - Be THOROUGH - try to find matches for EVERY item in our checklist
        - For each checklist item, try to find multiple places where changes could be made
        - Keep each individual change focused and surgical (don't rewrite entire paragraphs)

        RESPONSE FORMAT:
        Your response must be a valid JSON object with a "changes" array containing objects with these exact fields:
        - "original_text": The exactly copied text from the document with no modifications
        - "new_text": Focused replacement that implements our preferred language
        - "reason": Brief explanation referencing the specific checklist item

        TECHNICAL NOTE: Our text replacement system uses simple string matching. The text you provide as original_text will be searched for in the document and replaced with new_text only if it matches exactly. There is no fuzzy matching or pattern matching capability.
        """

        # Log the prompt being sent to OpenAI
        logger.info("Sending prompt to OpenAI for NDA analysis (JSON mode requested)")

        response = client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=[
                {
                    "role": "system",
                    "content": "You are a legal document reviewer specializing in NDAs. Your task is to analyze legal documents and suggest improvements based on the client's preferred language. Format your output as a JSON array of suggested changes.",
                },
                {"role": "user", "content": prompt},
            ],
            temperature=0.2,
            response_format={"type": "json_object"},
        )
        content = response.choices[0].message.content

        # Log the response content for debugging
        logger.info(
            f"OpenAI response received (JSON mode). Content length: {len(content)}"
        )
        logger.debug(f"Raw JSON response content: {content[:500]}...")

        # Parse the response
        try:
            # Parse the JSON response
            changes_raw = json.loads(content)
            logger.info(f"Successfully parsed JSON from OpenAI response")

            # Handle different possible JSON structures
            if isinstance(changes_raw, list):
                changes = changes_raw
            elif isinstance(changes_raw, dict) and isinstance(
                changes_raw.get("changes"), list
            ):
                changes = changes_raw["changes"]
            elif isinstance(changes_raw, dict) and isinstance(
                changes_raw.get("suggestions"), list
            ):
                changes = changes_raw["suggestions"]
            elif (
                isinstance(changes_raw, dict)
                and "original_text" in changes_raw
                and "new_text" in changes_raw
            ):
                # Handle case where a single change object is returned
                changes = [changes_raw]
            else:
                logger.error(f"Unexpected JSON structure received: {type(changes_raw)}")
                raise ValueError("Unexpected JSON structure received from OpenAI.")

            # Standardize field names if needed
            standardized_changes = []
            for change in changes:
                std_change = {}

                # Map field names to our standard format
                std_change["original_text"] = change.get("original_text") or change.get(
                    "original", ""
                )
                std_change["new_text"] = (
                    change.get("new_text")
                    or change.get("suggested_change")
                    or change.get("revised")
                    or change.get("new", "")
                )
                std_change["reason"] = change.get("reason") or change.get(
                    "justification", "No reason provided."
                )

                if std_change["original_text"] and std_change["new_text"]:
                    standardized_changes.append(std_change)

            # Return valid changes
            valid_changes = standardized_changes or changes
            logger.info(f"Returning {len(valid_changes)} valid changes.")
            return {"changes": valid_changes}

        except json.JSONDecodeError as e:
            logger.error(f"Failed to parse OpenAI response as JSON: {e}")
            return {"changes": [], "error": "Failed to parse JSON response from AI"}
        except ValueError as e:
            logger.error(f"Error processing JSON structure: {e}")
            return {"changes": [], "error": str(e)}

    except Exception as e:
        logger.error(f"Error in OpenAI analysis: {str(e)}", exc_info=True)
        raise NDAProcessingError(f"Failed to analyze NDA: {str(e)}")


def process_nda(docx_content: bytes) -> dict:
    """Process NDA document and return changes"""
    try:
        # Load the document
        doc = Document(BytesIO(docx_content))

        # Extract text from document
        nda_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

        # Analyze with OpenAI
        analysis_result = analyze_nda_with_openai(nda_text)

        if not analysis_result.get("changes"):
            logger.warning("No changes suggested by OpenAI")
            return {"error": "No changes were suggested for this document"}

        # Use the advanced document processor to apply changes and generate professional redlines
        redlined_doc = AdvancedDocumentProcessor.process_document_with_redlining(
            docx_content, analysis_result["changes"], author_tag="NDA Review"
        )

        # Generate a unique filename
        filename = f"redlined_nda_{uuid.uuid4().hex[:8]}.docx"

        # Store the document in memory
        processed_documents = application.config.get("processed_docs", {})
        processed_documents[filename] = redlined_doc  # Store the bytes directly

        return {
            "changes": analysis_result["changes"],
            "filename": filename,
            "document": BytesIO(redlined_doc),  # Return BytesIO for immediate use
        }

    except Exception as e:
        logger.error(f"Error processing NDA: {str(e)}")
        raise NDAProcessingError(f"Failed to process NDA: {str(e)}")


def verify_webhook_signature(request_data: str, signature: str) -> bool:
    """Verify the webhook signature from Brevo"""
    if not signature:
        logger.warning("No webhook signature provided")
        return False

    webhook_key = os.getenv("BREVO_WEBHOOK_KEY")
    if not webhook_key:
        logger.error("BREVO_WEBHOOK_KEY not set")
        return False

    # Create HMAC signature
    hmac_obj = hmac.new(
        webhook_key.encode("utf-8"), request_data.encode("utf-8"), hashlib.sha256
    )
    expected_signature = hmac_obj.hexdigest()

    # Compare signatures
    return hmac.compare_digest(signature, expected_signature)


@application.route("/")
def index():
    """Serve the main page"""
    return render_template("index.html")


@application.route("/process-nda", methods=["POST"])
def process_nda_route():
    """Handle NDA processing request"""
    try:
        if "file" not in request.files:
            logger.error("No file provided in request")
            return jsonify({"error": "No file provided"}), 400

        file = request.files["file"]
        if not file.filename.endswith(".docx"):
            logger.error(f"Invalid file type: {file.filename}")
            return jsonify({"error": "Please upload a .docx file"}), 400

        # Read the file content
        file_content = file.read()
        logger.info(f"Read file content, size: {len(file_content)} bytes")

        # Process the NDA
        result = process_nda(file_content)

        if "error" in result:
            logger.error(f"Error processing NDA: {result['error']}")
            return jsonify({"error": result["error"]}), 400

        # Initialize processed_docs if it doesn't exist
        if "processed_docs" not in application.config:
            application.config["processed_docs"] = {}
            logger.info("Initialized processed_docs in application config")

        # Store the document in memory
        filename = result["filename"]
        application.config["processed_docs"][filename] = result["document"].getvalue()
        logger.info(
            f"Stored document {filename} in memory, size: {len(application.config['processed_docs'][filename])} bytes"
        )

        # Return the changes and filename
        return jsonify({"changes": result["changes"], "filename": filename})

    except NDAProcessingError as e:
        logger.error(f"NDA processing error: {str(e)}")
        return jsonify({"error": str(e)}), 400
    except Exception as e:
        logger.error(f"Unexpected error in process_nda_route: {str(e)}", exc_info=True)
        return jsonify({"error": f"An unexpected error occurred: {str(e)}"}), 500


@application.route("/download/<filename>")
def download_file(filename):
    """Download the processed document"""
    try:
        # Log the requested filename
        logger.info(f"Attempting to download file: {filename}")

        # Check if we have any processed documents
        processed_docs = application.config.get("processed_docs", {})
        logger.info(f"Available processed documents: {list(processed_docs.keys())}")

        if filename not in processed_docs:
            logger.error(f"File not found in processed documents: {filename}")
            return jsonify({"error": "File not found"}), 404

        # Get the document bytes
        doc_bytes = processed_docs[filename]
        logger.info(f"Retrieved document bytes, size: {len(doc_bytes)}")

        if not doc_bytes:
            logger.error("Document bytes are empty")
            return jsonify({"error": "Document is empty"}), 500

        # Create a new BytesIO object
        doc_io = BytesIO(doc_bytes)
        doc_io.seek(0)  # Ensure we're at the start of the file

        logger.info("Sending file for download")
        return send_file(
            doc_io,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name=filename,
        )
    except Exception as e:
        logger.error(f"Error downloading file: {str(e)}", exc_info=True)
        return jsonify({"error": f"Failed to download file: {str(e)}"}), 500


if __name__ == "__main__":
    # Run as a local Flask application
    application.run(host="0.0.0.0", port=5001, debug=True)
