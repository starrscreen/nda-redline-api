import subprocess
import tempfile
import os
import platform
import logging
import zipfile
import tarfile
from pathlib import Path
from typing import Union, Tuple, Optional, List, Dict
from docx import Document
from io import BytesIO

__version__ = "0.0.4"  # Make sure this matches the version of binaries you have

logger = logging.getLogger(__name__)


class XmlPowerToolsEngine(object):
    """
    Uses external binary tools to create high-quality redlines.
    This is a more sophisticated approach than pure Python-based solutions.
    """

    def __init__(self, target_path: Optional[str] = None):
        self.target_path = target_path
        self.extracted_binaries_path = self.__get_binary_path()

    def __get_binary_path(self):
        # First check if the binary is directly available in bin directory
        base_path = os.path.dirname(os.path.abspath(__file__))
        bin_path = os.path.join(base_path, "bin")
        binary_name = self.__get_binary_name()
        direct_binary_path = os.path.join(bin_path, binary_name)

        if os.path.exists(direct_binary_path) and os.access(
            direct_binary_path, os.X_OK
        ):
            logger.info(f"Using pre-extracted binary at {direct_binary_path}")
            return direct_binary_path

        # If not found in bin directory, try Lambda Layer
        layer_path = "/opt/binaries"
        if os.path.exists(layer_path):
            full_binary_path = os.path.join(layer_path, binary_name)
            if os.path.exists(full_binary_path):
                return full_binary_path

        # Fall back to extracting from archive
        return self.__unzip_binary()

    def __unzip_binary(self):
        base_path = os.path.dirname(os.path.abspath(__file__))
        binaries_path = os.path.join(base_path, "binaries")
        target_path = (
            self.target_path if self.target_path else os.path.join(base_path, "bin")
        )

        if not os.path.exists(target_path):
            os.makedirs(target_path)

        binary_name = self.__get_binary_name()
        full_binary_path = os.path.join(target_path, binary_name)

        if not os.path.exists(full_binary_path):
            zip_path = os.path.join(binaries_path, self.__get_archive_name())
            logger.info(f"Extracting binary from {zip_path} to {target_path}")
            self.__extract_binary(zip_path, target_path)
            # Make the binary executable
            os.chmod(full_binary_path, 0o755)

        return full_binary_path

    def __extract_binary(self, zip_path: str, target_path: str):
        if zip_path.endswith(".zip"):
            with zipfile.ZipFile(zip_path, "r") as zip_ref:
                zip_ref.extractall(target_path)
        elif zip_path.endswith(".tar.gz"):
            with tarfile.open(zip_path, "r:gz") as tar_ref:
                tar_ref.extractall(target_path)

    def __get_binary_name(self):
        os_name = platform.system().lower()
        if os_name == "windows":
            return "redlines.exe"
        return "redlines"

    def __get_archive_name(self):
        os_name = platform.system().lower()
        arch = platform.machine().lower()

        if arch in ("x86_64", "amd64"):
            arch = "x64"
        elif arch in ("arm64", "aarch64"):
            arch = "arm64"
        else:
            raise EnvironmentError(f"Unsupported architecture: {arch}")

        if os_name == "linux":
            return f"linux-{arch}-{__version__}.tar.gz"
        elif os_name == "windows":
            return f"win-{arch}-{__version__}.zip"
        elif os_name == "darwin":
            return f"osx-{arch}-{__version__}.tar.gz"
        else:
            raise EnvironmentError("Unsupported OS")

    def run_redline(
        self,
        author_tag: str,
        original: Union[bytes, Path],
        modified: Union[bytes, Path],
    ) -> Tuple[bytes, Optional[str], Optional[str]]:
        temp_files = []
        try:
            target_path = tempfile.NamedTemporaryFile(delete=False).name
            original_path = (
                self._write_to_temp_file(original)
                if isinstance(original, bytes)
                else original
            )
            modified_path = (
                self._write_to_temp_file(modified)
                if isinstance(modified, bytes)
                else modified
            )
            temp_files.extend([target_path, original_path, modified_path])

            command = [
                self.extracted_binaries_path,
                author_tag,
                original_path,
                modified_path,
                target_path,
            ]

            try:
                result = subprocess.run(
                    command,
                    check=True,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    text=True,
                    timeout=60,  # Add 60-second timeout
                )
            except subprocess.TimeoutExpired:
                logger.error("Redline process timed out after 60 seconds")
                raise RuntimeError(
                    "Redline generation timed out. The document may be too large or complex."
                )

            stdout_output = (
                result.stdout
                if isinstance(result.stdout, str) and len(result.stdout) > 0
                else None
            )
            stderr_output = (
                result.stderr
                if isinstance(result.stderr, str) and len(result.stderr) > 0
                else None
            )

            redline_output = Path(target_path).read_bytes()

            return redline_output, stdout_output, stderr_output

        finally:
            self._cleanup_temp_files(temp_files)

    def _cleanup_temp_files(self, temp_files):
        for file_path in temp_files:
            try:
                os.remove(file_path)
            except OSError as e:
                print(f"Error deleting temp file {file_path}: {e}")

    def _write_to_temp_file(self, data):
        temp_file = tempfile.NamedTemporaryFile(delete=False)
        temp_file.write(data)
        temp_file.close()
        return temp_file.name


class AdvancedDocumentProcessor:
    """
    Combines NDA processing techniques from backend.py with binary-based redlining
    from redlines_old.py for optimal results.
    """

    @staticmethod
    def apply_changes_to_document(doc_content: bytes, changes: List[Dict]) -> bytes:
        """
        Apply the specified changes to a document while preserving formatting

        Args:
            doc_content: The content of the original document
            changes: List of changes with 'original_text' and 'new_text' keys

        Returns:
            BytesIO object with the modified document
        """
        try:
            # Load the document
            doc = Document(BytesIO(doc_content))

            # Add debugging to inspect the changes
            logger.info(f"Processing {len(changes)} changes")
            for i, change in enumerate(changes):
                logger.info(f"Change {i+1}: {change.keys()}")

            # Normalize field names - handle both original_text/new_text and other field names
            normalized_changes = []
            for change in changes:
                # Check for different possible field name combinations
                if "original_text" in change and "new_text" in change:
                    orig = change["original_text"]
                    new = change["new_text"]
                elif "original_text" in change and "suggested_change" in change:
                    orig = change["original_text"]
                    new = change["suggested_change"]
                elif "original" in change and "revised" in change:
                    orig = change["original"]
                    new = change["revised"]
                elif "original" in change and "new" in change:
                    orig = change["original"]
                    new = change["new"]
                else:
                    logger.warning(
                        f"Skipping change with unknown field names: {change.keys()}"
                    )
                    continue

                # Clean up the text - remove extra whitespace and normalize line endings
                orig = orig.strip().replace("\r\n", "\n")
                new = new.strip().replace("\r\n", "\n")

                # Debug the original and new text
                logger.info(f"Original text: {orig[:50]}...")
                logger.info(f"New text: {new[:50]}...")

                normalized_changes.append({"original_text": orig, "new_text": new})

            logger.info(f"Normalized {len(normalized_changes)} changes for processing")

            # Sort changes by length (longest first) to avoid partial replacements
            sorted_changes = sorted(
                normalized_changes, key=lambda x: len(x["original_text"]), reverse=True
            )

            # Track if any changes were made
            changes_made = False

            # Extract all text from the document to make finding exact matches easier
            all_paragraphs = []
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    all_paragraphs.append(
                        {
                            "type": "paragraph",
                            "index": i,
                            "text": para.text,
                            "object": para,
                        }
                    )

            # Extract table text
            for t_idx, table in enumerate(doc.tables):
                for r_idx, row in enumerate(table.rows):
                    for c_idx, cell in enumerate(row.cells):
                        for p_idx, para in enumerate(cell.paragraphs):
                            if para.text.strip():
                                all_paragraphs.append(
                                    {
                                        "type": "table_cell",
                                        "table_index": t_idx,
                                        "row_index": r_idx,
                                        "cell_index": c_idx,
                                        "para_index": p_idx,
                                        "text": para.text,
                                        "object": para,
                                    }
                                )

            logger.info(f"Extracted {len(all_paragraphs)} text elements for matching")

            # Find and apply exact matches
            for change in sorted_changes:
                original_text = change["original_text"]
                new_text = change["new_text"]
                found_match = False

                for para_info in all_paragraphs:
                    if original_text in para_info["text"]:
                        # Found a match - replace the text
                        paragraph = para_info["object"]
                        modified_text = para_info["text"].replace(
                            original_text, new_text
                        )

                        if modified_text != para_info["text"]:
                            # Clear and replace the text
                            paragraph.clear()
                            paragraph.add_run(modified_text)

                            # Update the text in our tracking structure for future matches
                            para_info["text"] = modified_text

                            changes_made = True
                            found_match = True

                            logger.info(
                                f"Applied change to {para_info['type']} - replaced: {original_text[:30]}... with: {new_text[:30]}..."
                            )

                if not found_match:
                    logger.warning(f"No match found for text: {original_text[:50]}...")

            if not changes_made:
                logger.warning("No changes were applied to the document")
            else:
                logger.info(f"Successfully applied changes to the document")

            # Save the modified document to memory
            output = BytesIO()
            doc.save(output)
            output.seek(0)
            return output.getvalue()

        except Exception as e:
            logger.error(f"Error applying changes to document: {str(e)}")
            raise

    @staticmethod
    def process_document_with_redlining(
        doc_content: bytes, changes: List[Dict], author_tag: str = "NDA Review"
    ) -> bytes:
        """
        Process a document by applying changes and creating a redlined version
        using the binary redlining engine for professional results.

        Args:
            doc_content: The content of the original document
            changes: List of changes to apply
            author_tag: Author tag for the redlining

        Returns:
            bytes containing the redlined document
        """
        try:
            # First apply changes to get the modified document
            modified_doc = AdvancedDocumentProcessor.apply_changes_to_document(
                doc_content, changes
            )

            # Create BytesIO objects for both documents
            original_io = BytesIO(doc_content)
            modified_io = BytesIO(modified_doc)

            # Use the XmlPowerToolsEngine to create the redlined version
            redline_engine = XmlPowerToolsEngine()
            redlined_doc, stdout, stderr = redline_engine.run_redline(
                author_tag=author_tag,
                original=original_io.getvalue(),
                modified=modified_io.getvalue(),
            )

            if stderr:
                logger.warning(f"Redline process warnings: {stderr}")

            return redlined_doc

        except Exception as e:
            logger.error(f"Error in document redlining process: {str(e)}")
            raise
