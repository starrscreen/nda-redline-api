from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from io import BytesIO
import tempfile
import os
import logging

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    Handles document processing operations including applying changes
    and generating redlined versions.
    """

    @staticmethod
    def apply_changes(docx_content, changes):
        """
        Apply a list of changes to a document.

        Args:
            docx_content (bytes): The original document content
            changes (list): List of change dictionaries with 'original_text' and 'new_text' keys

        Returns:
            BytesIO: A BytesIO object containing the modified document
        """
        try:
            # Load the original document
            doc = Document(BytesIO(docx_content))

            # Sort changes by length (longest first) to avoid partial replacements
            sorted_changes = sorted(
                changes, key=lambda x: len(x["original_text"]), reverse=True
            )

            changes_made = False

            # Process each paragraph
            for paragraph in doc.paragraphs:
                for change in sorted_changes:
                    if change["original_text"] in paragraph.text:
                        # Simple text replacement
                        paragraph.text = paragraph.text.replace(
                            change["original_text"], change["new_text"]
                        )
                        changes_made = True
                        logger.info(
                            f"Applied change: {change['original_text']} → {change['new_text']}"
                        )

            # Process tables if any
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for change in sorted_changes:
                                if change["original_text"] in paragraph.text:
                                    paragraph.text = paragraph.text.replace(
                                        change["original_text"], change["new_text"]
                                    )
                                    changes_made = True

            if not changes_made:
                logger.warning("No changes were applied to the document")

            # Save the modified document to BytesIO
            output = BytesIO()
            doc.save(output)
            output.seek(0)
            return output

        except Exception as e:
            logger.error(f"Error applying changes: {str(e)}")
            raise

    @staticmethod
    def create_redlined_document(original_docx, modified_docx):
        """
        Create a redlined version showing differences between original and modified documents.

        Args:
            original_docx (BytesIO): Original document
            modified_docx (BytesIO): Modified document

        Returns:
            BytesIO: A BytesIO object containing the redlined document
        """
        try:
            # Load both documents
            original_doc = Document(original_docx)
            modified_doc = Document(modified_docx)

            # Create a new document for the redlined version
            redlined_doc = Document()

            # Add styles and formatting from the original document
            for section in original_doc.sections:
                new_section = redlined_doc.add_section()
                new_section.page_height = section.page_height
                new_section.page_width = section.page_width
                new_section.left_margin = section.left_margin
                new_section.right_margin = section.right_margin
                new_section.top_margin = section.top_margin
                new_section.bottom_margin = section.bottom_margin

            # Collect paragraphs from both documents
            original_paragraphs = [p.text for p in original_doc.paragraphs]
            modified_paragraphs = [p.text for p in modified_doc.paragraphs]

            # Process paragraphs to identify changes
            for i, (orig_text, mod_text) in enumerate(
                zip(original_paragraphs, modified_paragraphs)
            ):
                paragraph = redlined_doc.add_paragraph()

                if orig_text != mod_text:
                    # This paragraph has changes - add with highlighting
                    run = paragraph.add_run(mod_text)
                    run.font.highlight_color = WD_COLOR_INDEX.RED
                else:
                    # No changes - just add the original text
                    paragraph.add_run(orig_text)

            # Handle case where modified document has more paragraphs
            if len(modified_paragraphs) > len(original_paragraphs):
                for i in range(len(original_paragraphs), len(modified_paragraphs)):
                    paragraph = redlined_doc.add_paragraph()
                    run = paragraph.add_run(modified_paragraphs[i])
                    run.font.highlight_color = WD_COLOR_INDEX.RED

            # Save the redlined document to BytesIO
            output = BytesIO()
            redlined_doc.save(output)
            output.seek(0)
            return output

        except Exception as e:
            logger.error(f"Error creating redlined document: {str(e)}")
            raise

    @staticmethod
    def process_document_with_changes(docx_content, changes):
        """
        Combined function to process a document with changes and create redlined version.

        Args:
            docx_content (bytes): The original document content
            changes (list): List of change dictionaries

        Returns:
            tuple: (modified_document, redlined_document) both as BytesIO objects
        """
        # Apply changes to get modified document
        original_doc_io = BytesIO(docx_content)
        original_doc_io.seek(0)

        modified_doc_io = DocumentProcessor.apply_changes(docx_content, changes)

        # Create redlined version
        original_doc_io.seek(0)
        modified_doc_io.seek(0)
        redlined_doc_io = DocumentProcessor.create_redlined_document(
            original_doc_io, modified_doc_io
        )

        # Reset file pointers
        modified_doc_io.seek(0)
        redlined_doc_io.seek(0)

        return modified_doc_io, redlined_doc_io
