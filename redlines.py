import subprocess
import tempfile
import os
import platform
import logging
import zipfile
import tarfile
from pathlib import Path
from typing import Union, Tuple, Optional
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX

__version__ = "0.0.4"  # Make sure this matches the version of binaries you have

logger = logging.getLogger(__name__)


class XmlPowerToolsEngine(object):
    def __init__(self, target_path: Optional[str] = None):
        self.target_path = target_path
        self.extracted_binaries_path = self.__get_binary_path()

    def __get_binary_path(self):
        # First check if the binary is directly available in bin directory
        base_path = os.path.dirname(os.path.abspath(__file__))
        bin_path = os.path.join(base_path, "bin")
        binary_name = self.__get_binary_name()
        direct_binary_path = os.path.join(bin_path, binary_name)

        if os.path.exists(direct_binary_path) and os.access(
            direct_binary_path, os.X_OK
        ):
            logger.info(f"Using pre-extracted binary at {direct_binary_path}")
            return direct_binary_path

        # If not found in bin directory, try Lambda Layer
        layer_path = "/opt/binaries"
        if os.path.exists(layer_path):
            full_binary_path = os.path.join(layer_path, binary_name)
            if os.path.exists(full_binary_path):
                return full_binary_path

        # Fall back to extracting from archive
        return self.__unzip_binary()

    def __unzip_binary(self):
        base_path = os.path.dirname(os.path.abspath(__file__))
        binaries_path = os.path.join(base_path, "binaries")
        target_path = (
            self.target_path if self.target_path else os.path.join(base_path, "bin")
        )

        if not os.path.exists(target_path):
            os.makedirs(target_path)

        binary_name = self.__get_binary_name()
        full_binary_path = os.path.join(target_path, binary_name)

        if not os.path.exists(full_binary_path):
            zip_path = os.path.join(binaries_path, self.__get_archive_name())
            logger.info(f"Extracting binary from {zip_path} to {target_path}")
            self.__extract_binary(zip_path, target_path)
            # Make the binary executable
            os.chmod(full_binary_path, 0o755)

        return full_binary_path

    def __extract_binary(self, zip_path: str, target_path: str):
        if zip_path.endswith(".zip"):
            with zipfile.ZipFile(zip_path, "r") as zip_ref:
                zip_ref.extractall(target_path)
        elif zip_path.endswith(".tar.gz"):
            with tarfile.open(zip_path, "r:gz") as tar_ref:
                tar_ref.extractall(target_path)

    def __get_binary_name(self):
        os_name = platform.system().lower()
        if os_name == "windows":
            return "redlines.exe"
        return "redlines"

    def __get_archive_name(self):
        os_name = platform.system().lower()
        arch = platform.machine().lower()

        if arch in ("x86_64", "amd64"):
            arch = "x64"
        elif arch in ("arm64", "aarch64"):
            arch = "arm64"
        else:
            raise EnvironmentError(f"Unsupported architecture: {arch}")

        if os_name == "linux":
            return f"linux-{arch}-{__version__}.tar.gz"
        elif os_name == "windows":
            return f"win-{arch}-{__version__}.zip"
        elif os_name == "darwin":
            return f"osx-{arch}-{__version__}.tar.gz"
        else:
            raise EnvironmentError("Unsupported OS")

    def run_redline(
        self,
        author_tag: str,
        original: Union[bytes, Path],
        modified: Union[bytes, Path],
    ) -> Tuple[bytes, Optional[str], Optional[str]]:
        temp_files = []
        try:
            target_path = tempfile.NamedTemporaryFile(delete=False).name
            original_path = (
                self._write_to_temp_file(original)
                if isinstance(original, bytes)
                else original
            )
            modified_path = (
                self._write_to_temp_file(modified)
                if isinstance(modified, bytes)
                else modified
            )
            temp_files.extend([target_path, original_path, modified_path])

            command = [
                self.extracted_binaries_path,
                author_tag,
                original_path,
                modified_path,
                target_path,
            ]

            try:
                result = subprocess.run(
                    command,
                    check=True,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    text=True,
                    timeout=60,  # Add 60-second timeout
                )
            except subprocess.TimeoutExpired:
                logger.error("Redline process timed out after 60 seconds")
                raise RuntimeError(
                    "Redline generation timed out. The document may be too large or complex."
                )

            stdout_output = (
                result.stdout
                if isinstance(result.stdout, str) and len(result.stdout) > 0
                else None
            )
            stderr_output = (
                result.stderr
                if isinstance(result.stderr, str) and len(result.stderr) > 0
                else None
            )

            redline_output = Path(target_path).read_bytes()

            return redline_output, stdout_output, stderr_output

        finally:
            self._cleanup_temp_files(temp_files)

    def _cleanup_temp_files(self, temp_files):
        for file_path in temp_files:
            try:
                os.remove(file_path)
            except OSError as e:
                print(f"Error deleting temp file {file_path}: {e}")

    def _write_to_temp_file(self, data):
        temp_file = tempfile.NamedTemporaryFile(delete=False)
        temp_file.write(data)
        temp_file.close()
        return temp_file.name


class Redlines:
    def __init__(self, input_path):
        self.doc = Document(input_path)

    def save(self, output_path):
        # Create a modified version with changes preserved from the original
        for paragraph in self.doc.paragraphs:
            self._process_text_element(paragraph)

        # Process each table
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._process_text_element(paragraph)

        # Save the document
        self.doc.save(output_path)

    def _process_text_element(self, paragraph):
        # Skip empty paragraphs
        if not paragraph.text.strip():
            return

        # Process each run in the paragraph individually to preserve formatting
        for run in paragraph.runs:
            if run.text.strip():
                # Keep original formatting but add red highlighting
                original_bold = run.bold
                original_italic = run.italic
                original_underline = run.underline
                original_font = run.font.name if run.font else None
                original_size = run.font.size if run.font else None
                original_color = run.font.color.rgb if run.font.color else None

                # Add red highlighting
                run.font.highlight_color = WD_COLOR_INDEX.RED

                # Preserve original formatting
                run.bold = original_bold
                run.italic = original_italic
                run.underline = original_underline
                if original_font:
                    run.font.name = original_font
                if original_size:
                    run.font.size = original_size
                if original_color:
                    run.font.color.rgb = original_color
